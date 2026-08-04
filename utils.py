# import os
# import re
# from docx import Document
# from docx.oxml.ns import qn
# from docx.table import Table
# from docx.text.paragraph import Paragraph

# DEFAULT_IMAGE_DIR = "static/question_images"


# def normalize_text(text):
#     if text is None:
#         return ""
#     return text.replace("\xa0", " ").strip()


# def extract_table_text(table):
#     rows = []
#     for row in table.rows:
#         cells = [normalize_text(cell.text) for cell in row.cells]
#         rows.append(" | ".join(cells))
#     return "\n".join(rows).strip()


# def save_image_from_run(run, output_dir, image_counter):
#     blip_elements = run._element.findall(
#         './/a:blip',
#         namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
#     )

#     if not blip_elements:
#         return None

#     r_id = blip_elements[0].get(qn('r:embed'))
#     if not r_id:
#         return None

#     image_part = run.part.related_parts[r_id]
#     image_data = image_part.blob

#     image_filename = f"question_image_{image_counter}.png"
#     image_path = os.path.join(output_dir, image_filename)

#     with open(image_path, 'wb') as f:
#         f.write(image_data)

#     return image_filename


# def iter_block_items(document):
#     body = document.element.body
#     for child in body.iterchildren():
#         if child.tag.endswith('p'):
#             yield Paragraph(child, document)
#         elif child.tag.endswith('tbl'):
#             yield Table(child, document)


# def append_with_newline(existing, new_text):
#     new_text = normalize_text(new_text)
#     if not new_text:
#         return existing
#     if not existing:
#         return new_text
#     return existing + "\n" + new_text


# def is_question_start(text):
#     return bool(re.match(r"^\d+[\.\)]\s*", text))


# def is_answer_line(text):
#     return bool(re.match(r"^(answer|correct answer|ans|answr)\s*:", text, re.IGNORECASE))


# def is_instruction_line(text):
#     lowered = text.lower().strip()
#     starters = (
#         "use the following",
#         "use matrices",
#         "use the information below",
#         "use the information to answer",
#         "refer to the following",
#         "answer question",
#         "answer questions",
#     )
#     return lowered.startswith(starters)


# def parse_answer_letter(text):
#     match = re.search(
#         r"^(?:answer|correct answer|ans|answr)\s*:\s*([A-Da-d])\b",
#         text.strip(),
#         re.IGNORECASE
#     )
#     return match.group(1).lower() if match else ""


# def get_marks_from_text(text):
#     match = re.search(r"\((\d+)\s*(?:mks|marks?)\)", text, re.IGNORECASE)
#     if match:
#         return int(match.group(1))
#     return 1


# def finalize_question(question, questions):
#     if question and normalize_text(question.get("question")):
#         questions.append(question)
#         return True
#     return False


# def parse_docx_questions(file_stream, image_output_dir=DEFAULT_IMAGE_DIR):
#     document = Document(file_stream)
#     questions = []

#     current_question = None
#     current_option = None
#     current_shared_block = ""
#     image_counter = 0
#     skipped = 0

#     os.makedirs(image_output_dir, exist_ok=True)

#     blocks = list(iter_block_items(document))
#     print(f"📄 Found {len(blocks)} total blocks")

#     def close_current_question():
#         nonlocal current_question, current_option, skipped
#         if current_question:
#             saved = finalize_question(current_question, questions)
#             if not saved:
#                 skipped += 1
#         current_question = None
#         current_option = None

#     for idx, block in enumerate(blocks):
#         print(f"\nProcessing block {idx}: {type(block).__name__}")

#         if isinstance(block, Paragraph):
#             text = normalize_text(block.text)

#             # Save images
#             for run in block.runs:
#                 image_name = save_image_from_run(run, image_output_dir, image_counter + 1)
#                 if image_name and current_question:
#                     image_counter += 1
#                     current_question["image"] = image_name
#                     print(f"📸 Saved image: {image_name}")

#             if not text:
#                 continue

#             print("TEXT:", text)

#             # Start shared instruction block
#             if is_instruction_line(text):
#                 close_current_question()
#                 current_shared_block = text
#                 current_option = None
#                 print("📘 Started shared block")
#                 continue

#             # New question
#             if is_question_start(text):
#                 close_current_question()

#                 full_question = text
#                 if current_shared_block:
#                     full_question = current_shared_block + "\n\n" + text
#                     current_shared_block = ""   # clear after first use

#                 current_question = {
#                     "question": full_question,
#                     "a": "",
#                     "b": "",
#                     "c": "",
#                     "d": "",
#                     "answer": "",
#                     "extra_content": None,
#                     "image": None,
#                     "marks": get_marks_from_text(text),
#                 }
#                 current_option = None
#                 print("🆕 New question started")
#                 continue

#             # Option
#             option_match = re.match(r"^\(?([A-Da-d])[\.\)]\s*(.*)$", text)
#             if option_match and current_question:
#                 label = option_match.group(1).lower()
#                 current_question[label] = text
#                 current_option = label
#                 print(f"🔠 Option {label.upper()}")
#                 continue

#             # Answer
#             if is_answer_line(text) and current_question:
#                 current_question["answer"] = parse_answer_letter(text)
#                 current_option = None
#                 print(f"✅ Answer: {current_question['answer'] or '(blank)'}")
#                 continue

#             # Continuation
#             if current_question:
#                 if current_option:
#                     current_question[current_option] = append_with_newline(
#                         current_question[current_option], text
#                     )
#                     print(f"↪ Continued option {current_option.upper()}")
#                 else:
#                     current_question["question"] = append_with_newline(
#                         current_question["question"], text
#                     )
#                     print("↪ Continued current question")
#             else:
#                 current_shared_block = append_with_newline(current_shared_block, text)
#                 print("↪ Added to shared block")

#         elif isinstance(block, Table):
#             table_text = extract_table_text(block)
#             if not table_text:
#                 continue

#             print("📊 Table detected")

#             if current_question:
#                 if current_option:
#                     current_question[current_option] = append_with_newline(
#                         current_question[current_option], table_text
#                     )
#                     print(f"↪ Table added to option {current_option.upper()}")
#                 else:
#                     current_question["question"] = append_with_newline(
#                         current_question["question"], table_text
#                     )
#                     print("↪ Table added to current question")
#             else:
#                 current_shared_block = append_with_newline(current_shared_block, table_text)
#                 print("↪ Table added to shared block")

#     close_current_question()

#     print(f"\n✅ Parsed {len(questions)} question(s)")
#     if skipped:
#         print(f"⚠️ Skipped {skipped} question(s)")

#     for i, q in enumerate(questions, 1):
#         print(f"\nQuestion {i}")
#         print(q["question"][:300])
#         print("A:", q["a"])
#         print("B:", q["b"])
#         print("C:", q["c"])
#         print("D:", q["d"])
#         print("Answer:", q["answer"] or "(blank)")

#     return questions


# def get_quiz_status(user_id):
#     return "active"


# def extract_drive_id(url):
#     patterns = [
#         r"https://drive\.google\.com/file/d/([A-Za-z0-9_-]+)",
#         r"https://drive\.google\.com/open\?id=([A-Za-z0-9_-]+)"
#     ]

#     for pattern in patterns:
#         match = re.search(pattern, url)
#         if match:
#             return match.group(1)

#     return url


# def get_drive_embed_url(drive_url_or_id):
#     file_id = extract_drive_id(drive_url_or_id)
#     return f"https://drive.google.com/file/d/{file_id}/preview"

import os
import re
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

DEFAULT_IMAGE_DIR = "static/question_images"


def normalize_text(text):
    """Clean up text from Word documents - preserving numbers and currency"""
    if text is None:
        return ""
    # Remove non-breaking spaces
    text = text.replace("\xa0", " ").replace("\u2022", "-")
    # Remove multiple spaces
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def extract_table_text(table):
    """Extract text from a Word table with proper formatting"""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # Get text from paragraphs in cell
            cell_text = "\n".join([normalize_text(p.text) for p in cell.paragraphs if p.text.strip()])
            cells.append(cell_text)
        rows.append(" | ".join(cells))
    return "\n".join(rows).strip()


def save_image_from_run(run, output_dir, image_counter):
    """Save an image from a Word run to disk"""
    try:
        blip_elements = run._element.findall(
            './/a:blip',
            namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
        )

        if not blip_elements:
            return None

        r_id = blip_elements[0].get(qn('r:embed'))
        if not r_id:
            return None

        image_part = run.part.related_parts[r_id]
        image_data = image_part.blob

        # Determine image extension from content type
        ext = 'png'
        if hasattr(image_part, 'content_type'):
            if 'jpeg' in image_part.content_type or 'jpg' in image_part.content_type:
                ext = 'jpg'
            elif 'gif' in image_part.content_type:
                ext = 'gif'
            elif 'png' in image_part.content_type:
                ext = 'png'

        image_filename = f"question_image_{image_counter}.{ext}"
        image_path = os.path.join(output_dir, image_filename)

        with open(image_path, 'wb') as f:
            f.write(image_data)

        return image_filename
    except Exception as e:
        print(f"⚠️ Error saving image: {e}")
        return None


def iter_block_items(document):
    """Iterate through paragraphs and tables in a document"""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith('p'):
            yield Paragraph(child, document)
        elif child.tag.endswith('tbl'):
            yield Table(child, document)


def append_with_newline(existing, new_text):
    """Append text with proper newline formatting"""
    new_text = normalize_text(new_text)
    if not new_text:
        return existing
    if not existing:
        return new_text
    return existing + "\n" + new_text


def is_question_start(text):
    """Check if text starts a new question - handles numbered questions"""
    text = text.strip()
    # Match patterns like: 1. 1) 1. 1) 1. 1) 1. 1)
    # Also handles "16." for double digit numbers
    return bool(re.match(r"^(\d+)[\.\)]\s*", text))


def is_option_line(text):
    """Check if text is an option (A, B, C, D)"""
    text = text.strip()
    # Match patterns like: A. B) C. D)
    # Handle cases where option is followed by text
    return bool(re.match(r"^([A-Da-d])[\.\)]\s*", text))


def is_answer_line(text):
    """Check if text is an answer line"""
    text = text.strip().lower()
    return bool(re.match(r"^(answer|correct answer|ans)\s*[:.]?\s*([A-Da-d])", text))


def extract_answer(text):
    """Extract the answer letter from text"""
    text = text.strip().lower()
    match = re.search(r"(answer|correct answer|ans)\s*[:.]?\s*([A-Da-d])", text)
    if match:
        return match.group(2).lower()
    return ""


def get_marks_from_text(text):
    """Extract marks from question text"""
    text = text.strip()
    # Match patterns like: (2 marks), (2mks), (2)
    match = re.search(r"\((\d+)\s*(?:mks|marks?)?\)", text, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return 1  # Default marks


def is_financial_table(table_text):
    """Check if table contains financial data"""
    table_text = table_text.lower()
    financial_terms = [
        'shs', 'debit', 'credit', 'balance', 'account', 'amount',
        'revenue', 'expense', 'asset', 'liability', 'equity',
        'income', 'profit', 'loss', 'cash', 'bank', 'capital',
        'opening', 'closing', 'total', 'subtotal', 'cost', 'selling',
        'sales', 'purchases', 'inventory', 'stock', 'receivable',
        'payable', 'loan', 'interest', 'tax', 'dividend', 'depreciation'
    ]
    count = sum(1 for term in financial_terms if term in table_text)
    return count >= 3


def is_financial_data(text):
    """Check if text contains financial data"""
    text = text.lower().strip()
    # Check for currency symbols or financial terms
    financial_patterns = [
        r'(?:shs|kshs|kes)\s*[\d,]+',  # Shs 000
        r'[\$\€\£\₦\₹]',  # Currency symbols
        r'\d{1,3}(?:,\d{3})*(?:\.\d{2})?',  # Numbers with commas
        r'debit|credit|dr|cr',
        r'balance sheet|income statement|trial balance',
        r'profit|loss|revenue|expense|asset|liability|equity',
        r'cost|selling|inventory|stock|receivable|payable|loan',
        r'interest|tax|dividend|depreciation'
    ]
    return any(re.search(pattern, text, re.IGNORECASE) for pattern in financial_patterns)


def is_instruction_line(text):
    """Check if text contains exam instructions - expanded for accounting"""
    text = text.lower().strip()
    instruction_keywords = [
        "use the following",
        "use matrices",
        "use the information",
        "refer to the following",
        "answer question",
        "answer questions",
        "read the following",
        "study the following",
        "consider the following",
        "based on the following",
        "using the following",
        "given the following",
        "the following table",
        "the following information",
        "the following data",
        "the following scenario",
        "the following passage",
        "the following text",
        "the following case",
        "the following diagram",
        "the following figure",
        "the following chart",
        "the following graph",
        "the following trial balance",
        "the following balance sheet",
        "the following income statement",
        "the following financial statements",
        "the following accounts",
        "the following transactions",
        "the following ledger",
        "the following journal",
        "the following inventory",
        "the following vehicles",
        "the following assets",
        "the following liabilities",
        "the following statements",
        "statement of profit or loss",
        "statement of financial position",
        "use the following information to answer",
        "prepare the following",
        "calculate the following",
        "compute the following",
        "determine the following",
        "analyze the following",
        "the following additional information",
        "the following further information"
    ]
    return any(keyword in text for keyword in instruction_keywords)


def parse_docx_questions(file_path, image_output_dir=DEFAULT_IMAGE_DIR):
    """
    Parse questions from a DOCX file with enhanced support for accounting questions
    """
    try:
        document = Document(file_path)
    except Exception as e:
        print(f"❌ Error opening document: {e}")
        return []

    questions = []
    current_question = None
    current_option = None
    shared_instructions = []
    image_counter = 0
    skipped = 0
    in_instruction_block = False
    instruction_counter = 0

    os.makedirs(image_output_dir, exist_ok=True)

    # Get all blocks (paragraphs and tables)
    blocks = list(iter_block_items(document))
    print(f"📄 Found {len(blocks)} blocks to process")

    def save_current_question():
        """Save the current question if it's valid"""
        nonlocal current_question, skipped
        
        if not current_question:
            return
        
        # Check if question has at least some content
        question_text = normalize_text(current_question.get("question", ""))
        if question_text:
            questions.append(current_question)
            print(f"✅ Saved question {len(questions)}")
        else:
            skipped += 1

    def clean_shared_instructions():
        """Convert shared instructions to a single string"""
        nonlocal shared_instructions
        if shared_instructions:
            result = "\n\n".join(shared_instructions)
            shared_instructions = []
            return result
        return ""

    # Process each block
    for idx, block in enumerate(blocks):
        if isinstance(block, Paragraph):
            text = normalize_text(block.text)
            if not text:
                continue

            print(f"\n📝 Block {idx}: {text[:100]}...")

            # Check for images in the paragraph
            for run in block.runs:
                if current_question and not current_question.get("image"):
                    image_name = save_image_from_run(run, image_output_dir, image_counter + 1)
                    if image_name:
                        image_counter += 1
                        current_question["image"] = image_name
                        print(f"📸 Saved image: {image_name}")

            # Check if this is an instruction line
            if not current_question and is_instruction_line(text):
                shared_instructions.append(text)
                in_instruction_block = True
                instruction_counter += 1
                print(f"📘 Added to shared instructions ({instruction_counter}): {text[:50]}...")
                continue

            # Check if this starts a new question
            if is_question_start(text):
                # Save previous question
                save_current_question()

                # Get shared instructions if any
                instructions = clean_shared_instructions()
                
                # Create new question with shared instructions if any
                full_text = text
                if instructions:
                    full_text = instructions + "\n\n" + text
                    in_instruction_block = False

                current_question = {
                    "question": full_text,
                    "a": "",
                    "b": "",
                    "c": "",
                    "d": "",
                    "answer": "",
                    "image": None,
                    "marks": get_marks_from_text(text),
                    "extra_content": ""
                }
                current_option = None
                print(f"🆕 New question started: {text[:50]}...")
                continue

            # Check if this is an option
            option_match = re.match(r"^([A-Da-d])[\.\)]\s*(.*)", text)
            if option_match and current_question:
                label = option_match.group(1).lower()
                option_text = option_match.group(2).strip()
                current_question[label] = option_text
                current_option = label
                print(f"🔠 Option {label.upper()}: {option_text[:50]}...")
                continue

            # Check if this is an answer line
            if is_answer_line(text) and current_question:
                current_question["answer"] = extract_answer(text)
                print(f"✅ Answer: {current_question['answer'] or 'Not found'}")
                continue

            # If we have a current question, append text to it
            if current_question:
                if current_option:
                    # Append to current option
                    current_question[current_option] = append_with_newline(
                        current_question[current_option], text
                    )
                    print(f"↪ Added to option {current_option.upper()}")
                else:
                    # Check if this is financial data
                    if is_financial_data(text):
                        text = "📊 " + text
                    
                    # Append to question text
                    current_question["question"] = append_with_newline(
                        current_question["question"], text
                    )
                    print(f"↪ Added to question")
            else:
                # No current question, add to shared instructions
                shared_instructions.append(text)
                print(f"↪ Added to shared instructions")

        elif isinstance(block, Table):
            table_text = extract_table_text(block)
            if not table_text:
                continue

            print(f"📊 Table found: {table_text[:100]}...")

            # Format financial tables nicely
            if is_financial_table(table_text):
                table_text = "┌─────────────────────────────────────┐\n" + \
                             "📊 FINANCIAL DATA:\n" + \
                             table_text + \
                             "\n└─────────────────────────────────────┘"

            if current_question:
                if current_option:
                    current_question[current_option] = append_with_newline(
                        current_question[current_option], table_text
                    )
                    print(f"↪ Table added to option {current_option.upper()}")
                else:
                    current_question["question"] = append_with_newline(
                        current_question["question"], table_text
                    )
                    print(f"↪ Table added to question")
            else:
                shared_instructions.append(table_text)
                print(f"↪ Table added to shared instructions")

    # Save the last question
    save_current_question()

    print(f"\n✅ Parsed {len(questions)} questions successfully")
    if skipped:
        print(f"⚠️ Skipped {skipped} invalid questions")

    # Debug output for first few questions
    for i, q in enumerate(questions[:3], 1):
        print(f"\n📋 Question {i}:")
        print(f"   Text: {q['question'][:100]}...")
        if q.get('a'):
            print(f"   A: {q['a'][:50]}")
        if q.get('b'):
            print(f"   B: {q['b'][:50]}")
        if q.get('c'):
            print(f"   C: {q['c'][:50]}")
        if q.get('d'):
            print(f"   D: {q['d'][:50]}")
        print(f"   Answer: {q['answer']}")
        if q.get('image'):
            print(f"   Image: {q['image']}")

    return questions


def get_quiz_status(user_id):
    """Get the status of a quiz for a user"""
    return "active"


def extract_drive_id(url):
    """Extract Google Drive file ID from URL"""
    patterns = [
        r"https://drive\.google\.com/file/d/([A-Za-z0-9_-]+)",
        r"https://drive\.google\.com/open\?id=([A-Za-z0-9_-]+)",
        r"https://drive\.google\.com/uc\?id=([A-Za-z0-9_-]+)",
    ]

    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)

    return url


def get_drive_embed_url(drive_url_or_id):
    """Convert Google Drive URL to embed URL"""
    file_id = extract_drive_id(drive_url_or_id)
    return f"https://drive.google.com/file/d/{file_id}/preview"